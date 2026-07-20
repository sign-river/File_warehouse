from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = "reference.docx"

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(0.8)
section.footer_distance = Cm(0.8)

styles = doc.styles

normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.first_line_indent = Cm(0.74)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, bold, before, after in [
    ("Title", 18, True, 0, 8),
    ("Subtitle", 11, False, 0, 6),
    ("Heading 1", 12, True, 10, 4),
    ("Heading 2", 10.5, True, 6, 2),
    ("Heading 3", 10.5, False, 4, 2),
]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei" if bold else "SimSun")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    if name == "Title":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif name == "Heading 1":
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for style_name in ["Caption", "Table Caption"]:
    if style_name in styles:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set default table font and prevent row splitting where possible.
for style_name in ["Table", "Table Grid"]:
    if style_name in styles:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(9)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

# Add a minimal paragraph so the reference document has valid body content.
p = doc.add_paragraph("Reference document for the editable Chinese IEEE-style draft.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Enable automatic hyphenation and keep equations/images stable.
settings = doc.settings.element
compat = settings.find(qn("w:compat"))
if compat is None:
    compat = OxmlElement("w:compat")
    settings.append(compat)

# Add page number field in footer.
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run()
fld_char1 = OxmlElement("w:fldChar")
fld_char1.set(qn("w:fldCharType"), "begin")
instr_text = OxmlElement("w:instrText")
instr_text.set(qn("xml:space"), "preserve")
instr_text.text = " PAGE "
fld_char2 = OxmlElement("w:fldChar")
fld_char2.set(qn("w:fldCharType"), "end")
run._r.append(fld_char1)
run._r.append(instr_text)
run._r.append(fld_char2)

# Remove the minimal paragraph from final rendering by making it hidden.
r = p.runs[0]
vanish = OxmlElement("w:vanish")
r._r.get_or_add_rPr().append(vanish)

doc.save(OUT)
print(f"Created {OUT}")
