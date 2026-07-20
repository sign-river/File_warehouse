from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = "reference.docx"

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.9)
section.left_margin = Cm(1.9)
section.right_margin = Cm(1.9)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

for name, size, bold in [
    ("Title", 18, True),
    ("Subtitle", 11, False),
    ("Heading 1", 12, True),
    ("Heading 2", 10.5, True),
    ("Heading 3", 10.5, False),
]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")

styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 1"].paragraph_format.space_before = Pt(8)
styles["Heading 1"].paragraph_format.space_after = Pt(4)
styles["Heading 2"].paragraph_format.space_before = Pt(6)
styles["Heading 2"].paragraph_format.space_after = Pt(3)

# Create a minimal body so Pandoc can inherit stable styles.
p = doc.add_paragraph("IEEE-style editable draft", style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("正文示例。", style="Normal")

doc.save(OUT)
print(f"Created {OUT}")
